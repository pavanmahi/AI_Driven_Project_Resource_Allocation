# from fastapi import APIRouter, HTTPException, UploadFile, File
# from typing import List, Optional
# from pydantic import BaseModel, Field
# from datetime import datetime, date
# from pathlib import Path
# import json
# import pdfplumber
# from docx import Document
# from ..services.llm_service import ProjectRequirementExtractor

# router = APIRouter()
# extractor = ProjectRequirementExtractor()

# DATA_FILE = Path(__file__).resolve().parent.parent.parent / "data" / "projects.json"

# # ---------- Models ----------
# class ProjectBase(BaseModel):
#     name: str
#     description: str
#     phase: str = "Design"
#     status: str = "Planning"
#     team_size_required: int = 3
#     priority_level: str = Field(default="Medium", description="High, Medium, or Low")
#     budget: Optional[float] = None
#     start_date: Optional[date] = None
#     end_date: Optional[date] = None
#     timeline: Optional[List[str]] = []

# class ProjectCreate(ProjectBase):
#     pass

# class Project(ProjectBase):
#     id: int
#     created_at: datetime
#     updated_at: datetime


# # ---------- Helpers ----------
# def load_projects() -> List[Project]:
#     if not DATA_FILE.exists():
#         return []
#     with open(DATA_FILE, "r", encoding="utf-8") as f:
#         data = json.load(f)
#     return [Project(**p) for p in data]

# def save_projects(projects: List[Project]):
#     with open(DATA_FILE, "w", encoding="utf-8") as f:
#         json.dump([p.dict() for p in projects], f, indent=4, default=str)

# def extract_text_from_file(file: UploadFile) -> str:
#     """Extract text from PDF or DOCX UploadFile"""
#     filename = file.filename.lower()

#     if filename.endswith(".pdf"):
#         text = ""
#         with pdfplumber.open(file.file) as pdf:
#             for page in pdf.pages:
#                 text += page.extract_text() or ""
#         return text.strip()

#     elif filename.endswith(".docx"):
#         doc = Document(file.file)
#         return "\n".join([para.text for para in doc.paragraphs]).strip()

#     return "Unsupported file type"


# # ---------- Routes ----------
# @router.get("/", response_model=List[Project])
# async def get_projects():
#     """Get all projects"""
#     return load_projects()

# @router.post("/create", response_model=Project)
# async def create_project(project: ProjectCreate):
#     """Create a new project manually"""
#     projects = load_projects()
#     new_project = Project(
#         id=(max([p.id for p in projects], default=0) + 1),
#         **project.dict(),
#         created_at=datetime.utcnow(),
#         updated_at=datetime.utcnow()
#     )
#     projects.append(new_project)
#     save_projects(projects)
#     return new_project

# @router.post("/upload", response_model=Project)
# async def create_project_from_file(file: UploadFile = File(...)):
#     """Create a new project from uploaded file (PDF/DOCX)"""
#     text = extract_text_from_file(file)
#     description = await extractor.generate_short_description(text)

#     projects = load_projects()
#     new_project = Project(
#         id=(max([p.id for p in projects], default=0) + 1),
#         name=file.filename,
#         description=description,
#         created_at=datetime.utcnow(),
#         updated_at=datetime.utcnow()
#     )
#     projects.append(new_project)
#     save_projects(projects)
#     return new_project

# @router.get("/{project_id}", response_model=Project)
# async def get_project(project_id: int):
#     """Get a project by ID"""
#     projects = load_projects()
#     project = next((p for p in projects if p.id == project_id), None)
#     if not project:
#         raise HTTPException(status_code=404, detail="Project not found")
#     return project

# @router.put("/{project_id}", response_model=Project)
# async def update_project(project_id: int, project: ProjectCreate):
#     """Update a project"""
#     projects = load_projects()
#     index = next((i for i, p in enumerate(projects) if p.id == project_id), None)
#     if index is None:
#         raise HTTPException(status_code=404, detail="Project not found")

#     updated_project = Project(
#         id=project_id,
#         **project.dict(),
#         created_at=projects[index].created_at,
#         updated_at=datetime.utcnow()
#     )
#     projects[index] = updated_project
#     save_projects(projects)
#     return updated_project

# @router.delete("/{project_id}")
# async def delete_project(project_id: int):
#     """Delete a project"""
#     projects = load_projects()
#     index = next((i for i, p in enumerate(projects) if p.id == project_id), None)
#     if index is None:
#         raise HTTPException(status_code=404, detail="Project not found")

#     deleted_project = projects.pop(index)
#     save_projects(projects)
#     return {"message": f"Project '{deleted_project.name}' deleted successfully"}
from fastapi import APIRouter, HTTPException, UploadFile, File
from typing import List, Optional
from pydantic import BaseModel, Field
from datetime import datetime, date
from pathlib import Path
import json
import pdfplumber
from docx import Document
from ...llm_service import ProjectRequirementExtractor

router = APIRouter()
extractor = ProjectRequirementExtractor()

DATA_FILE = Path(__file__).resolve().parent.parent.parent / "data" / "projects.json"

# ---------- Models ----------
class ProjectBase(BaseModel):
    name: str
    description: str
    phase: str = "Design"
    status: str = "Planning"
    team_size_required: int = 3
    priority_level: str = Field(default="Medium", description="High, Medium, or Low")
    budget: Optional[float] = None
    start_date: Optional[date] = None
    end_date: Optional[date] = None
    timeline: Optional[List[str]] = []

class ProjectCreate(ProjectBase):
    pass

from typing import Optional, Union
from pydantic import BaseModel

class Project(BaseModel):
    id: Optional[int] = None
    name: Optional[str] = None
    priority_level: Optional[Union[str, int]] = None
    created_at: Optional[str] = None
    updated_at: Optional[str] = None
    project_name: Optional[str] = None
    deadline: Optional[str] = None


# ---------- Helpers ----------
def load_projects() -> List[Project]:
    if not DATA_FILE.exists():
        return []
    with open(DATA_FILE, "r", encoding="utf-8") as f:
        data = json.load(f)
    return [Project(**p) for p in data]

def save_projects(projects: List[Project]):
    with open(DATA_FILE, "w", encoding="utf-8") as f:
        json.dump([p.dict() for p in projects], f, indent=4, default=str)

def extract_text_from_file(file: UploadFile) -> str:
    """Extract text from PDF or DOCX UploadFile"""
    filename = file.filename.lower()

    if filename.endswith(".pdf"):
        text = ""
        with pdfplumber.open(file.file) as pdf:
            for page in pdf.pages:
                text += page.extract_text() or ""
        return text.strip()

    elif filename.endswith(".docx"):
        doc = Document(file.file)
        return "\n".join([para.text for para in doc.paragraphs]).strip()

    return "Unsupported file type"


# ---------- Routes ----------
@router.get("/", response_model=List[Project])
async def get_projects():
    """Get all projects"""
    return load_projects()

@router.post("/create", response_model=Project)
async def create_project(project: ProjectCreate):
    """Create a new project manually"""
    projects = load_projects()
    new_project = Project(
        id=(max([p.id for p in projects], default=0) + 1),
        **project.dict(),
        created_at=datetime.utcnow(),
        updated_at=datetime.utcnow()
    )
    projects.append(new_project)
    save_projects(projects)
    return new_project

@router.post("/upload", response_model=Project)
async def create_project_from_file(file: UploadFile = File(...)):
    """Create a new project by uploading a document"""
    text = extract_text_from_file(file)   # already synchronous

    # Generate project description using LLM (sync call)
    description = extractor.generate_short_description(text)

    # Load existing projects
    projects = load_projects()

    new_project = Project(
        id=(max([p.id for p in projects], default=0) + 1),
        name=file.filename,
        description=description,
        status="Planning",
        created_at=datetime.utcnow(),
        updated_at=datetime.utcnow()
    )

    projects.append(new_project)
    save_projects(projects)

    return new_project


@router.get("/{project_id}", response_model=Project)
async def get_project(project_id: Optional[int] = None):
    projects = load_projects()
    if project_id is None:
        return projects
    """Get a project by ID"""
    project = next((p for p in projects if p.id == project_id), None)
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    return project

@router.put("/{project_id}", response_model=Project)
async def update_project(project_id: int, project: ProjectCreate):
    """Update a project"""
    projects = load_projects()
    index = next((i for i, p in enumerate(projects) if p.id == project_id), None)
    if index is None:
        raise HTTPException(status_code=404, detail="Project not found")

    updated_project = Project(
        id=project_id,
        **project.dict(),
        created_at=projects[index].created_at,
        updated_at=datetime.utcnow()
    )
    projects[index] = updated_project
    save_projects(projects)
    return updated_project

@router.delete("/{project_id}")
async def delete_project(project_id: int):
    """Delete a project"""
    projects = load_projects()
    index = next((i for i, p in enumerate(projects) if p.id == project_id), None)
    if index is None:
        raise HTTPException(status_code=404, detail="Project not found")

    deleted_project = projects.pop(index)
    save_projects(projects)
    return {"message": f"Project '{deleted_project.name}' deleted successfully"}
